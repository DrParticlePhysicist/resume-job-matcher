# # File: backend/services.py

# import spacy
# from sklearn.feature_extraction.text import TfidfVectorizer
# from sklearn.metrics.pairwise import cosine_similarity
# import fitz  # PyMuPDF for PDF processing
# from docx import Document
# from typing import List

# # Load spaCy model
# nlp = spacy.load("en_core_web_sm")

# def preprocess_text(text: str) -> str:
#     """Cleans and preprocesses text."""
#     doc = nlp(text)
#     return " ".join([token.lemma_ for token in doc if not token.is_stop and not token.is_punct])

# def extract_text_from_pdf(pdf_file) -> str:
#     """Extracts text from a PDF file."""
#     doc = fitz.open(stream=pdf_file.file.read(), filetype="pdf")
#     text = " ".join([page.get_text("text") for page in doc])
#     return text

# def extract_text_from_docx(docx_file) -> str:
#     """Extracts text from a DOCX file."""
#     doc = Document(docx_file.file)
#     return " ".join([para.text for para in doc.paragraphs])

# def process_resume(file) -> str:
#     """Processes an uploaded resume and extracts text."""
#     if file.content_type == "application/pdf":
#         return extract_text_from_pdf(file)
#     elif file.content_type == "application/msword":
#         return extract_text_from_docx(file)
#     elif file.content_type == "text/plain":
#         return file.file.read().decode("utf-8")
#     else:
#         raise ValueError("Unsupported file format")

# def match_job_descriptions(resume_text: str, job_descriptions: List[str]):
#     """Matches resume text against multiple job descriptions using cosine similarity."""
#     vectorizer = TfidfVectorizer()
#     texts = [resume_text] + job_descriptions
#     tfidf_matrix = vectorizer.fit_transform(texts)
#     similarities = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:])
    
#     # Return scores sorted with job descriptions
#     scored_jobs = sorted(zip(job_descriptions, similarities[0]), key=lambda x: x[1], reverse=True)
#     return scored_jobs
import pandas as pd
import numpy as np
import spacy
import fitz  # PyMuPDF for PDF processing
from docx import Document
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# Load spaCy model
nlp = spacy.load("en_core_web_sm")

# Load Kaggle dataset (used for weight boosting)
def load_kaggle_skills():
    kaggle_path = "../data/kaggle_dataset.csv"
    df = pd.read_csv(kaggle_path)
    
    # Extract all unique required skills from the dataset
    kaggle_skills = set()
    for skills in df["required_skills"].dropna():
        kaggle_skills.update(skills.split(", "))
    
    return kaggle_skills

# Load skills for weight boosting
KAGGLE_SKILLS = load_kaggle_skills()

def preprocess_text(text: str) -> str:
    """Cleans and preprocesses text."""
    doc = nlp(text)
    return " ".join([token.lemma_ for token in doc if not token.is_stop and not token.is_punct])

def extract_text_from_pdf(pdf_file) -> str:
    """Extracts text from a PDF file."""
    doc = fitz.open(stream=pdf_file.file.read(), filetype="pdf")
    text = " ".join([page.get_text("text") for page in doc])
    return text

def extract_text_from_docx(docx_file) -> str:
    """Extracts text from a DOCX file."""
    doc = Document(docx_file.file)
    return " ".join([para.text for para in doc.paragraphs])

def process_resume(file) -> str:
    """Processes an uploaded resume and extracts text."""
    if file.content_type == "application/pdf":
        return extract_text_from_pdf(file)
    elif file.content_type == "application/msword":
        return extract_text_from_docx(file)
    elif file.content_type == "text/plain":
        return file.file.read().decode("utf-8")
    else:
        raise ValueError("Unsupported file format")

def compute_match_score(resume_text, job_descriptions):
    """Matches resume text against multiple job descriptions using cosine similarity with Kaggle-based skill boosting."""
    resume_text = preprocess_text(resume_text)
    job_descriptions = [preprocess_text(job) for job in job_descriptions]
    
    documents = [resume_text] + job_descriptions
    vectorizer = TfidfVectorizer()
    tfidf_matrix = vectorizer.fit_transform(documents)
    
    # Extract feature names (words)
    feature_names = vectorizer.get_feature_names_out()
    
    # Convert TF-IDF matrix to a DataFrame for easy manipulation
    tfidf_df = pd.DataFrame(tfidf_matrix.toarray(), columns=feature_names)
    
    # Apply weight boosting for Kaggle-listed skills
    boost_factor = 1.5  # Increase weight for Kaggle-listed words
    for skill in KAGGLE_SKILLS:
        if skill in feature_names:
            tfidf_df[skill] *= boost_factor

    # Convert back to NumPy array after modification
    boosted_tfidf_matrix = tfidf_df.to_numpy()
    
    # Compute Cosine Similarity
    resume_vector = boosted_tfidf_matrix[0].reshape(1, -1)  # First row is the resume
    job_vectors = boosted_tfidf_matrix[1:]  # Remaining rows are job descriptions
    similarity_scores = cosine_similarity(resume_vector, job_vectors)[0]

    # Pair job descriptions with their similarity scores
    matched_results = list(zip(job_descriptions, similarity_scores))  # [(job_desc1, score1), (job_desc2, score2), ...]

    # Sort in descending order of similarity
    matched_results.sort(key=lambda x: x[1], reverse=True)

    return matched_results  # Returns a list of (job_desc, score) tuples
